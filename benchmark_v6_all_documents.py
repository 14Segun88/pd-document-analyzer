#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для прогона всех документов из папки "Анализ пакета 1" через LLM Mistral v6
С сохранением результатов и эталонов в MD файлы для сравнения

Запуск:
    cd /home/segun/CascadeProjects/Перед\ 0_2
    python3 benchmark_v6_all_documents.py
"""

import os
import sys
import json
import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Optional

# Добавляем путь к v6
sys.path.insert(0, str(Path(__file__).parent / "v6"))

try:
    from web_app_v6_cot_fallback import DocumentAnalyzer
except ImportError as e:
    print(f"❌ Ошибка импорта: {e}")
    print("Убедитесь, что файл v6/web_app_v6_cot_fallback.py существует")
    sys.exit(1)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


def load_etalons_from_docx(docx_path: Path) -> Dict[str, Dict]:
    """Загружает эталоны из DOCX файла"""
    try:
        from docx import Document
        doc = Document(str(docx_path))
        
        etalons = {}
        current_doc = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text and (text.startswith('1.') or text.startswith('2.') or 
                        text.startswith('3.') or text.startswith('4.')):
                # Новый документ
                current_doc = text.split('.')[1].strip().split('(')[0].strip()
                etalons[current_doc] = {
                    'title': None,
                    'customer': None,
                    'developer': None,
                    'year': None,
                    'document_type': None,
                    'content_summary': None,
                    'purpose': None
                }
        
        # Извлекаем данные из таблиц
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    key = cells[0].lower()
                    value = cells[1]
                    
                    # Маппинг ключей
                    if 'назван' in key:
                        etalons.setdefault('current', {})['title'] = value
                    elif 'заказчик' in key:
                        etalons.setdefault('current', {})['customer'] = value
                    elif 'подрядчик' in key or 'разработчик' in key:
                        etalons.setdefault('current', {})['developer'] = value
                    elif 'год' in key:
                        etalons.setdefault('current', {})['year'] = value
                    elif 'тип' in key and 'документ' in key:
                        etalons.setdefault('current', {})['document_type'] = value
                    elif 'содерж' in key:
                        etalons.setdefault('current', {})['content_summary'] = value
                    elif 'цель' in key:
                        etalons.setdefault('current', {})['purpose'] = value
        
        return etalons
    except Exception as e:
        logger.error(f"Ошибка загрузки эталонов: {e}")
        return {}


def find_etalon_in_kb(filename: str, kb_data: List[Dict]) -> Optional[Dict]:
    """Находит эталон в KB по имени файла"""
    import re
    
    # Извлекаем шифр проекта и раздел
    file_section = re.search(r'ПД[№\s]*(\d+)', filename)
    file_code = re.search(r'(МКД-\d{4}-\d{4}|МЕС-БМК-\d{2}/\d{2}|МЕС_БМК_\d{2}_\d{2}|\d{2,4}/\d{2,4})', filename)
    
    best_match = None
    best_score = 0
    
    for entry in kb_data:
        kb_title = entry.get('title', '')
        
        # Проверяем совпадение по шифру
        if file_code:
            file_code_norm = file_code.group(1).replace('_', '-')
            if file_code_norm in kb_title:
                # Проверяем раздел
                if file_section:
                    section = file_section.group(1)
                    if f'Раздел {section}' in kb_title or f'Раздел{section}' in kb_title:
                        best_match = entry
                        best_score = 1.0
                        break
    
    # Если не нашли по шифру, ищем по другим признакам
    if not best_match:
        for entry in kb_data:
            kb_title = entry.get('title', '').lower()
            filename_lower = filename.lower()
            
            # Проверяем тип документа
            if 'ЕГРН' in filename or 'Выписка' in filename:
                if 'егрн' in kb_title or 'выписка' in kb_title:
                    best_match = entry
                    break
            elif 'ГПЗУ' in filename:
                if 'гпзу' in kb_title:
                    best_match = entry
                    break
            elif 'Доверен' in filename:
                if 'доверен' in kb_title:
                    best_match = entry
                    break
            elif 'Программа' in filename:
                if 'программа' in kb_title:
                    best_match = entry
                    break
            elif 'ТЗ' in filename:
                if 'тз' in kb_title or 'задание' in kb_title:
                    best_match = entry
                    break
            elif 'ТУ' in filename or 'Технические условия' in filename:
                if 'ту' in kb_title or 'условия' in kb_title:
                    best_match = entry
                    break
    
    return best_match


def generate_md_report(filename: str, llm_result: Dict, etalon: Optional[Dict], 
                       timestamp: str, output_dir: Path) -> str:
    """Генерирует MD отчёт для одного документа"""
    
    md_content = f"""# Анализ документа: {filename}

**Дата анализа:** {timestamp}
**Модель:** Mistral 14B Reasoning (v6_cot_fallback)

---

## 📊 Результат LLM

| Поле | Значение |
|------|----------|
| **Название** | {llm_result.get('title', '—')} |
| **Заказчик** | {llm_result.get('customer', '—')} |
| **Подрядчик / проектная фирма** | {llm_result.get('developer', '—')} |
| **Год составления** | {llm_result.get('year', '—')} |
| **Тип документа** | {llm_result.get('document_type', '—')} |
| **Содержимое** | {llm_result.get('content_summary', '—')} |
| **Цель** | {llm_result.get('purpose', '—')} |

---

## 📋 Эталон из KB

"""
    
    if etalon:
        md_content += f"""| Поле | Значение |
|------|----------|
| **Название** | {etalon.get('title', '—')} |
| **Заказчик** | {etalon.get('customer', '—')} |
| **Подрядчик / проектная фирма** | {etalon.get('developer', '—')} |
| **Год составления** | {etalon.get('year', '—')} |
| **Тип документа** | {etalon.get('document_type', '—')} |
| **Содержимое** | {etalon.get('content_summary', '—')} |
| **Цель** | {etalon.get('purpose', '—')} |

---
"""
    else:
        md_content += "**❌ Эталон не найден в KB**\n\n---\n"
    
    # Сравнение
    md_content += "\n## 🔍 Сравнение LLM vs Эталон\n\n"
    
    if etalon:
        fields = ['title', 'customer', 'developer', 'year', 'document_type', 'content_summary', 'purpose']
        
        for field in fields:
            llm_val = llm_result.get(field, '—')
            etalon_val = etalon.get(field, '—')
            
            # Нормализация для сравнения
            llm_norm = str(llm_val).lower().strip() if llm_val else ''
            etalon_norm = str(etalon_val).lower().strip() if etalon_val else ''
            
            if llm_norm == etalon_norm and llm_norm:
                status = "✅ Совпадает"
            elif not etalon_norm or etalon_val == '—':
                status = "⚠️ Нет эталона"
            elif llm_norm in etalon_norm or etalon_norm in llm_norm:
                status = "⚠️ Частичное"
            else:
                status = "❌ Не совпадает"
            
            md_content += f"**{field}:** {status}\n"
            md_content += f"- LLM: {llm_val}\n"
            md_content += f"- Эталон: {etalon_val}\n\n"
    else:
        md_content += "**Сравнение невозможно: эталон не найден**\n"
    
    return md_content


def main():
    """Главная функция"""
    
    # Пути
    input_dir = Path("/home/segun/CascadeProjects/Перед 0_2/Перед 0/Isxodnie_documenti/Анализ пакета 1")
    output_dir = Path("/home/segun/CascadeProjects/Перед 0_2/Тесты_md")
    kb_path = Path("/home/segun/CascadeProjects/Перед 0_2/knowledge_base.json")
    
    # Создаём выходную директорию
    output_dir.mkdir(exist_ok=True)
    
    # Загружаем KB
    with open(kb_path, 'r', encoding='utf-8') as f:
        kb_data = json.load(f)
    
    logger.info(f"KB loaded: {len(kb_data)} entries")
    
    # Инициализируем анализатор
    analyzer = DocumentAnalyzer()
    
    # Получаем список файлов
    all_files = []
    for f in input_dir.iterdir():
        if f.suffix.lower() in ['.pdf', '.xml']:
            all_files.append(f)
    
    logger.info(f"Найдено файлов: {len(all_files)}")
    
    # Прогоняем каждый файл
    timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    
    results = []
    
    for i, filepath in enumerate(sorted(all_files), 1):
        filename = filepath.name
        
        logger.info(f"\n{'='*80}")
        logger.info(f"[{i}/{len(all_files)}] Обработка: {filename}")
        logger.info(f"{'='*80}")
        
        try:
            # Анализируем документ
            if filepath.suffix.lower() == '.pdf':
                result = analyzer.analyze_pdf(filepath, filename)
            elif filepath.suffix.lower() == '.xml':
                result = analyzer.analyze_xml(filepath, filename)
            else:
                continue
            
            # Находим эталон в KB
            etalon = find_etalon_in_kb(filename, kb_data)
            
            # Генерируем MD отчёт
            md_content = generate_md_report(
                filename=filename,
                llm_result=result,
                etalon=etalon,
                timestamp=timestamp,
                output_dir=output_dir
            )
            
            # Сохраняем MD файл
            safe_filename = filename.replace('/', '_').replace('\\', '_').replace(' ', '_')
            md_filename = f"analysis_{timestamp}_{safe_filename}.md"
            md_path = output_dir / md_filename
            
            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(md_content)
            
            logger.info(f"✅ Сохранено: {md_filename}")
            
            results.append({
                'filename': filename,
                'llm_result': result,
                'etalon': etalon,
                'md_path': str(md_path)
            })
            
        except Exception as e:
            logger.error(f"❌ Ошибка обработки {filename}: {e}")
            results.append({
                'filename': filename,
                'error': str(e)
            })
    
    # Генерируем сводный отчёт
    summary_path = output_dir / f"SUMMARY_{timestamp}.md"
    
    with open(summary_path, 'w', encoding='utf-8') as f:
        f.write(f"# Сводный отчёт: {len(results)} документов\n\n")
        f.write(f"**Дата:** {timestamp}\n\n")
        f.write(f"**Модель:** Mistral 14B Reasoning (v6_cot_fallback)\n\n")
        f.write("---\n\n")
        
        for result in results:
            filename = result['filename']
            if 'error' in result:
                f.write(f"❌ **{filename}** - Ошибка: {result['error']}\n\n")
            else:
                md_path = Path(result['md_path']).name
                f.write(f"✅ **{filename}** - [{md_path}]({md_path})\n\n")
    
    logger.info(f"\n{'='*80}")
    logger.info(f"Завершено! Обработано: {len(results)} документов")
    logger.info(f"Сводный отчёт: {summary_path}")
    logger.info(f"{'='*80}")


if __name__ == '__main__':
    main()
